"""Resume text extraction service."""
from typing import Optional
import re
from pathlib import Path
from io import BytesIO


def extract_text_from_pdf(file_path: Optional[str] = None, file_bytes: Optional[bytes] = None) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader

        if file_bytes:
            reader = PdfReader(BytesIO(file_bytes))
        elif file_path:
            reader = PdfReader(file_path)
        else:
            return ""

        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
        return text.strip()
    except Exception as e:
        return f"[PDF extraction error: {e}]"


def extract_text_from_docx(file_path: Optional[str] = None, file_bytes: Optional[bytes] = None) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document

        if file_bytes:
            doc = Document(BytesIO(file_bytes))
        elif file_path:
            doc = Document(file_path)
        else:
            return ""

        text = "\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
        return text.strip()
    except Exception as e:
        return f"[DOCX extraction error: {e}]"


def extract_text_from_latex(file_bytes: Optional[bytes] = None, file_path: Optional[str] = None) -> str:
    """Extract plain text from a LaTeX file by stripping commands."""
    try:
        if file_bytes:
            text = file_bytes.decode("utf-8", errors="replace")
        elif file_path:
            text = Path(file_path).read_text(errors="replace")
        else:
            return ""

        # Remove comments
        text = re.sub(r'%.*$', '', text, flags=re.MULTILINE)
        # Remove \begin{...} and \end{...}
        text = re.sub(r'\\(begin|end)\{[^}]*\}', '', text)
        # Remove \command{...} but keep content inside braces
        text = re.sub(r'\\[a-zA-Z]+\{([^}]*)\}', r'\1', text)
        # Remove remaining \commands
        text = re.sub(r'\\[a-zA-Z]+\*?', '', text)
        # Remove braces
        text = re.sub(r'[{}]', '', text)
        # Clean up whitespace
        text = re.sub(r'\n{3,}', '\n\n', text)
        return text.strip()
    except Exception as e:
        return f"[LaTeX extraction error: {e}]"


def extract_resume_text(filename: str, file_path: Optional[str] = None, file_bytes: Optional[bytes] = None) -> str:
    """Extract text from a resume file based on extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_path=file_path, file_bytes=file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_path=file_path, file_bytes=file_bytes)
    elif ext == ".tex":
        return extract_text_from_latex(file_bytes=file_bytes, file_path=file_path)
    elif ext == ".txt":
        if file_bytes:
            return file_bytes.decode("utf-8", errors="replace")
        elif file_path:
            return Path(file_path).read_text(errors="replace")
    return ""


def parse_contact_info(text: str) -> dict:
    """Try to extract name, email, phone from resume text."""
    result = {"name": "", "email": "", "phone": ""}

    email_match = re.search(r'[\w.+-]+@[\w-]+\.[\w.-]+', text)
    if email_match:
        result["email"] = email_match.group()

    phone_match = re.search(r'[\+]?[\d\s\-\(\)]{10,15}', text)
    if phone_match:
        result["phone"] = phone_match.group().strip()

    lines = text.strip().split('\n')
    for line in lines[:5]:
        line = line.strip()
        if line and not re.search(r'[@\d]', line) and len(line) < 60:
            result["name"] = line
            break

    return result


# Lines/phrases that strongly indicate a job-description document rather
# than a candidate's resume. Tuned conservatively — we only want to fire
# when several of these co-occur near the top of the document.
_JD_PHRASES = re.compile(
    r"\b("
    r"job\s+description|position\s+description|role\s+description|"
    r"about\s+the\s+role|about\s+the\s+position|about\s+the\s+job|"
    r"key\s+responsibilities|primary\s+responsibilities|"
    r"what\s+you[’']ll\s+do|what\s+we[’']re\s+looking\s+for|"
    r"minimum\s+qualifications|preferred\s+qualifications|"
    r"required\s+qualifications|nice\s+to\s+have|"
    r"reports?\s+to|reporting\s+to|"
    r"we\s+are\s+looking\s+for|we\s+are\s+hiring|"
    r"compensation\s+(?:and|&)\s+benefits|equal\s+opportunity\s+employer|"
    r"company\s+overview|about\s+(?:us|the\s+company)"
    r")\b",
    re.IGNORECASE,
)

_RESUME_PHRASES = re.compile(
    r"\b("
    r"work\s+experience|professional\s+experience|employment\s+history|"
    r"education\s*(?:and|&)?\s*(?:certifications)?\s*[:\n]|"
    r"skills?\s*[:\n]|technical\s+skills|core\s+competencies|"
    r"certifications?\s*[:\n]|projects?\s*[:\n]|"
    r"references\s+available|curriculum\s+vitae|résumé|resume\s*[:\n]"
    r")\b",
    re.IGNORECASE,
)


def looks_like_job_description(text: str, filename: str = "") -> bool:
    """Best-effort detection of an uploaded *job description* — i.e. the
    user dropped a JD PDF into the candidate uploader by mistake.

    Returns True when either:
      • the filename strongly signals a JD (``jd_*``, ``*_jd.pdf``,
        ``job_description*``), OR
      • the text has 3+ JD-style section headers near the top AND no
        candidate-shaped sections (Work Experience / Skills / Education).

    Conservative on purpose — false positives block a real CV upload, so
    we only fire on clear cases.
    """
    fname_lower = (filename or "").strip().lower()
    if fname_lower:
        stem = re.sub(r"\.[a-z0-9]{1,5}$", "", fname_lower)
        if (
            stem.startswith("jd_") or stem.startswith("jd-") or
            stem.endswith("_jd") or stem.endswith("-jd") or
            "job_description" in stem or "job-description" in stem or
            stem.startswith("position_") or stem.startswith("role_")
        ):
            return True

    if not text:
        return False

    head = "\n".join(text.split("\n")[:80])
    jd_hits = len(set(m.group(0).lower() for m in _JD_PHRASES.finditer(head)))
    resume_hits = len(set(m.group(0).lower() for m in _RESUME_PHRASES.finditer(text)))

    # Strong JD signal: many JD headers, very few/no resume sections, and
    # no email anywhere in the first 80 lines (real candidates always
    # include contact details near the top).
    has_email_in_head = bool(re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", head))
    if jd_hits >= 3 and resume_hits <= 1 and not has_email_in_head:
        return True
    return False
